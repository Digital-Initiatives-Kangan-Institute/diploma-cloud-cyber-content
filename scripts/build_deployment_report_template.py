#!/usr/bin/env python3
"""Build the YAT / MTS Deployment Report template (.docx) from the brand pack.

ONE generic, in-world Deployment Report template — the superset that serves any
deployment, from a greenfield foundation build to a full HA hardening. Built from the
HA report outline (the larger of the two), with AT2's foundation-only sections folded
back in, and assessment-only scaffolding (KE Q&A, reflections, student ID, UoC tags)
stripped — the same generic-template principle applied to the Business Case.

Sections a simpler deployment doesn't need carry an explicit "Not applicable — reason"
convention rather than being deleted, so the report always proves completeness.

Usage:  python scripts/build_deployment_report_template.py [output.docx]
Default: ../diploma-cloud-cyber-website/public/templates/YAT-Deployment-Report-template.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_bc_template as bc  # noqa: E402  shared brand helpers

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def applic(doc, hint):
    """Small terracotta applicability marker for conditional (e.g. HA) sections."""
    p = doc.add_paragraph()
    r = p.add_run(f"Applicability: complete for {hint}; otherwise mark this section "
                  "“Not applicable — [reason]”.")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(bc.TERRACOTTA)
    p.paragraph_format.space_after = Pt(6)
    return p


def callout(doc, lines):
    """Shaded convention box."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    bc.shade_cell(cell, bc.CREAM)
    bc.set_cell_borders(cell, bc.OCHRE, sz=8)
    for i, (head, body) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        hr = p.add_run(head)
        hr.bold = True
        hr.font.size = Pt(10)
        hr.font.color.rgb = RGBColor.from_string(bc.TEAL)
        if body:
            br = p.add_run("  " + body)
            br.font.size = Pt(10)
            br.font.color.rgb = RGBColor.from_string(bc.CHARCOAL)
    doc.add_paragraph()


def build(path):
    doc = Document()
    bc.configure_styles(doc)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    bc.build_header_footer(sec)

    # ---- COVER ----
    bc.wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(bc.ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(bc.GREY)
    bc.paragraph_bottom_rule(doc.add_paragraph(), bc.TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Deployment Report")
    sub = doc.add_paragraph().add_run("[ Deployment / engagement name ]")
    sub.font.size = Pt(15); sub.italic = True; sub.font.color.rgb = RGBColor.from_string(bc.GREY)
    doc.add_paragraph()

    cover_rows = [
        ("Engagement", "[ Engagement name ]"),
        ("Engagement reference", "[ Reference ]"),
        ("Report version", "[ e.g. v1.0 ]"),
        ("Prepared by", "[ Consultant name ]"),
        ("Consultant role", "[ Role, e.g. MTS Consultant ]"),
        ("Date submitted", "[ DD/MM/YYYY ]"),
        ("Submitted to", "[ Acceptance authority / sponsor ]"),
        ("Related documents", "[ e.g. the approved design this report implements ]"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    ct.alignment = WD_TABLE_ALIGNMENT.LEFT
    for k, v in cover_rows:
        cells = ct.add_row().cells
        bc.set_cell_borders(cells[0]); bc.set_cell_borders(cells[1]); bc.shade_cell(cells[0], bc.CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        vr = cells[1].paragraphs[0].add_run(v); vr.font.size = Pt(10); vr.italic = True
        vr.font.color.rgb = RGBColor.from_string(bc.GREY)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS + convention ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    doc.add_paragraph("How to use this template", style="Heading 1")
    callout(doc, [
        ("One report for any deployment.", "This is the standard YAT deployment report — it covers "
         "everything from a greenfield foundation build to a full high-availability hardening."),
        ("Complete every section.", "Where a section does not apply to your deployment, mark it "
         "“Not applicable — [reason]” rather than deleting it, so the report proves "
         "nothing was overlooked."),
        ("Sections flagged with an applicability note", "(in terracotta) are the ones a simpler "
         "deployment will often mark Not applicable — e.g. the HA failure/resize simulations."),
        ("Cross-reference your evidence.", "The build narrative and testing sections reference the "
         "screenshots and configuration exports captured in the appendices."),
    ])
    doc.add_paragraph("Contents", style="Heading 1")
    bc.add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h2 = lambda t: doc.add_paragraph(t, style="Heading 2")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    h1("1. Executive Summary")
    bc.guidance(doc, "Write this last. A ≤ 1-page summary the reader sees first: what was deployed (or "
                     "changed); the region/AZ footprint; the 2–3 highlights; any limitations or items "
                     "deferred to a later phase; and — for a resilience deployment — whether the "
                     "availability target is now met and the headline simulation outcomes. ~250–400 words.")
    bc.response(doc)

    h1("2. Engagement Context")
    bc.guidance(doc, "Brief context for the reader (≤ ½ page): the strategic/prior work this deployment "
                     "builds on (the approved business case and the design being implemented), your role, "
                     "and the scope hand-off to any later phase.")
    bc.response(doc)

    h1("3. Scope of Deployment")
    bc.guidance(doc, "What is included in this deployment and what is deferred (≤ ½ page). Restate from the "
                     "approved design in your own words: the in-scope components, and what is out of scope / "
                     "deferred to a later phase.")
    bc.response(doc)
    h3("3.1 Maintenance-window context")
    applic(doc, "a change to a running production system")
    bc.guidance(doc, "If this work was performed on a live system within a maintenance window, describe the "
                     "window (duration, timing) and the requirement to finish either with the work complete "
                     "or cleanly rolled back. For a greenfield build, mark Not applicable.")
    bc.response(doc)

    h1("4. Build / Change Narrative")
    bc.guidance(doc, "A layer-by-layer account of what was built or changed. For each layer, write a short "
                     "narrative — for a fresh build, what you stood up; for a change to an existing build, "
                     "what changed from the baseline — and cross-reference the Appendix A screenshots and "
                     "Appendix B configuration exports.")
    for n, title, hint in [
        ("4.1", "Identity and access management (IAM)", "account access, groups/users/roles, MFA, instance profiles — or, for a change, any IAM additions (state if none)"),
        ("4.2", "Network topology", "VPC, subnets, gateways, route tables — or the subnets/routes added (e.g. a second AZ)"),
        ("4.3", "Compute (EC2 + Auto Scaling)", "launch template, instance type + why, ASG min/desired/max + scaling policy — or capacity/AZ-spread changes"),
        ("4.4", "Load balancing (ALB)", "the load balancer, target group, listener, health check — or cross-AZ target changes"),
        ("4.5", "Database (RDS)", "instance class + why, engine version, storage, encryption, backups — or enabling Multi-AZ"),
        ("4.6", "Storage (EBS + S3)", "EBS volumes, S3 buckets, encryption, public-access settings"),
        ("4.7", "Security (security groups + encryption)", "the tiered security-group model, encryption in transit + at rest — or HA-related adjustments"),
        ("4.8", "Monitoring", "the CloudWatch alarms and thresholds — baseline, or the HA-tuned set (per-AZ counts, failover detection, replica lag, service-level dashboard)"),
    ]:
        h3(f"{n} {title}")
        bc.guidance(doc, f"Cover: {hint}. Cross-reference the relevant Appendix A screenshots and Appendix B exports.")
        bc.response(doc)
    h3("4.9 Cross-Region backup / replication")
    applic(doc, "deployments with cross-Region resilience in scope")
    bc.guidance(doc, "Cross-Region snapshot copies, AWS Backup, or S3 cross-Region replication, if in scope. "
                     "Otherwise mark Not applicable.")
    bc.response(doc)

    h1("5. Configuration Decisions")
    bc.guidance(doc, "The approved design leaves specific decisions to the implementer. For each decision you "
                     "made, state your choice and justify it against the workload and requirements. Add rows "
                     "as needed; examples below span foundation-build and HA decisions.")
    bc.table(doc, ["#", "Decision point", "Your decision", "Rationale"],
             [["C1", "EC2 instance type (within the chosen family)", "[ … ]", "[ workload + cost ]"],
              ["C2", "RDS instance class", "[ … ]", "[ … ]"],
              ["C3", "Storage sizing (EBS data volume / RDS storage)", "[ show the calc ]", "[ data footprint + growth ]"],
              ["C4", "ASG scaling threshold", "[ … ]", "[ expected CPU profile ]"],
              ["C5", "RDS Multi-AZ apply timing (HA)", "[ … ]", "[ maintenance-window rationale ]"],
              ["C6", "Post-HA ASG capacity (min ≥ 2 across AZs)", "[ … ]", "[ AZ-failure resilience ]"],
              ["C7", "Cross-Region backup destination (HA)", "[ … ]", "[ … ]"],
              ["…", "[ add further decisions ]", "[ … ]", "[ … ]"]],
             widths=[1.0, 6.0, 4.0, 4.5])

    h1("6. Testing, Simulation and Validation")
    bc.guidance(doc, "Document the tests run to verify the deployment. For each test, state the test, the "
                     "result, and reference the supporting evidence in Appendix C.")
    h3("6.1 Connectivity tests")
    bc.table(doc, ["Test", "Outcome (Pass/Fail)", "Notes"],
             [["ALB → EC2 health check", "[ ]", "[ ]"],
              ["EC2 → RDS connection (private)", "[ ]", "[ ]"],
              ["EC2 → internet via NAT", "[ ]", "[ ]"],
              ["RDS not publicly reachable (negative test)", "[ ]", "[ ]"]],
             widths=[7.0, 4.0, 4.5])
    h3("6.2 Autoscaling test")
    bc.guidance(doc, "Trigger a scaling event (e.g. load against the ALB) and confirm the ASG scales out and "
                     "back in. Cross-reference Appendix C.")
    bc.response(doc)
    h3("6.3 Database connectivity and basic operations")
    bc.guidance(doc, "Confirm the database tier is reachable from the app tier over the private network, the "
                     "engine version meets the application requirement, and encryption-in-transit is in place.")
    bc.response(doc)
    h3("6.4 Infrastructure end-to-end smoke test")
    bc.guidance(doc, "Confirm the infrastructure is ready to serve traffic (e.g. a placeholder page via the "
                     "ALB DNS returns HTTP 200/302 from a backend instance that can reach the database).")
    bc.response(doc)
    h3("6.5 Failure simulation")
    applic(doc, "high-availability / resilience deployments")
    bc.guidance(doc, "Execute each failure simulation from the design (e.g. EC2 termination, RDS Multi-AZ "
                     "failover, AZ partition). Record method, observed outcome, and whether the service "
                     "stayed reachable. Otherwise mark Not applicable.")
    bc.table(doc, ["#", "Simulation", "Method", "Observed outcome", "Reachable throughout?", "Evidence"],
             [["F1", "[ e.g. EC2 termination ]", "[ … ]", "[ … ]", "[ Yes/No ]", "C—"],
              ["F2", "[ e.g. RDS failover ]", "[ … ]", "[ … ]", "[ Yes (brief blip) ]", "C—"]],
             widths=[0.8, 3.0, 3.0, 3.6, 2.6, 1.5])
    h3("6.6 Resize simulation")
    applic(doc, "high-availability / resilience deployments")
    bc.guidance(doc, "Execute each resize simulation from the design (e.g. ASG capacity increase, RDS instance "
                     "class change) and record the availability impact. Otherwise mark Not applicable.")
    bc.response(doc)
    h3("6.7 Availability measurement")
    applic(doc, "deployments with an availability target to demonstrate")
    bc.guidance(doc, "Describe how availability was measured (e.g. CloudWatch dashboard, a curl-loop against "
                     "the ALB during the window) and report the measured availability across the window.")
    bc.response(doc)
    h3("6.8 Simulation findings vs the design")
    applic(doc, "high-availability / resilience deployments")
    bc.guidance(doc, "For each simulation, compare the observed outcome against the expected outcome in the "
                     "design; document any divergence and why.")
    bc.response(doc)
    h3("6.9 Adjustments made per simulation outcomes")
    applic(doc, "high-availability / resilience deployments")
    bc.guidance(doc, "Any changes made to the architecture, configuration, or monitoring based on what the "
                     "simulations revealed — or an explicit statement that none were needed, with reasoning.")
    bc.response(doc)

    h1("7. Operational Handover")
    bc.guidance(doc, "Hand-over information for the team taking over the infrastructure.")
    h3("7.1 Access")
    bc.guidance(doc, "Who has what access post-handover, MFA enforcement, any IAM group changes.")
    bc.response(doc)
    h3("7.2 Runbook references")
    bc.guidance(doc, "Pointers to the design document, naming/tagging conventions, backup arrangements, and "
                     "the alarms list + notification destinations.")
    bc.response(doc)
    h3("7.3 Known limitations and what's next")
    bc.guidance(doc, "Be explicit about what is not covered today and what a later phase would add.")
    bc.response(doc)
    h3("7.4 Documentation filing")
    bc.table(doc, ["Item", "Filed in", "Reference"],
             [["This Deployment Report", "[ YAT ICT shared documentation ]", "[ ref ]"],
              ["Configuration exports (Appendix B)", "[ … ]", "[ ref ]"],
              ["Test / simulation evidence (Appendix C)", "[ … ]", "[ ref ]"]],
             widths=[6.5, 5.0, 4.0])
    h3("7.5 Feedback record")
    applic(doc, "deployments where stakeholder feedback is captured")
    bc.table(doc, ["Feedback received", "From", "Your response", "Resulting action"],
             [["[ … ]", "[ … ]", "[ … ]", "[ … ]"]],
             widths=[5.0, 3.0, 4.0, 4.0])
    h3("7.6 Sign-off")
    bc.table(doc, ["Role", "Name", "Date", "Signature"],
             [["Prepared by", "[ … ]", "", ""],
              ["Reviewed by", "[ … ]", "", ""],
              ["Approved by (acceptance authority)", "[ … ]", "", ""]],
             widths=[5.5, 4.5, 2.5, 3.0])

    # ---- APPENDICES ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    h1("Appendix A — Build evidence (screenshots)")
    bc.guidance(doc, "Capture a console screenshot evidencing each component you built or changed, with the "
                     "region indicator visible. List each below and cross-reference it from §4 / §6. Examples:")
    bc.table(doc, ["#", "Screenshot", "What must be visible"],
             [["A1", "IAM groups / MFA", "[ created groups; MFA enabled ]"],
              ["A2", "VPC subnets", "[ subnets + AZs ]"],
              ["A3", "EC2 instances + ASG", "[ running instances across AZs; ASG min/desired/max ]"],
              ["A4", "ALB target group health", "[ healthy targets ]"],
              ["A5", "RDS database", "[ available; Multi-AZ status; encryption ]"],
              ["A6", "CloudWatch alarms / dashboard", "[ the alarms / service-level dashboard ]"],
              ["…", "[ add as your deployment requires ]", "[ … ]"]],
             widths=[1.0, 5.5, 9.0])
    h1("Appendix B — Configuration exports")
    bc.guidance(doc, "Export each configuration (AWS CLI or console) and attach as a code block or file. "
                     "Examples: IAM policies; security-group rules; VPC/subnet/route tables; launch template "
                     "+ ASG; ALB + target groups; RDS instance; S3 bucket policy/encryption; CloudWatch alarms.")
    bc.response(doc, "[ Configuration exports ]")
    h1("Appendix C — Test and simulation evidence")
    bc.guidance(doc, "Attach the evidence supporting the results in §6 — screenshots, terminal/log excerpts, "
                     "metric graphs, and (for HA work) failure/resize simulation captures and the computed "
                     "availability over the window.")
    bc.response(doc, "[ Test and simulation evidence ]")

    h1("Document control")
    bc.table(doc, ["Field", "Value"],
             [["Document version", "[ v1.0 ]"],
              ["Author", "[ Name, role ]"],
              ["Engagement", "[ Engagement name ]"],
              ["Date submitted", "[ DD/MM/YYYY ]"],
              ["Distribution", "[ … ]"],
              ["Related documents", "[ the design implemented; predecessor/successor reports ]"]],
             widths=[5.0, 10.5])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "../diploma-cloud-cyber-website/public/templates/YAT-Deployment-Report-Template.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
