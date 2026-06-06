#!/usr/bin/env python3
"""Build the AT1 Solution Design EXEMPLAR (.docx) — assessor marking reference.

A complete, worked model answer for the YAT LMS Global Expansion Solution Design (S1-CL2 AT1
Part A, ICTCLD503 *design*). It mirrors the YAT Solution Design template's section structure
(build_solution_design_template.py) so the exemplar and the student template stay in parity,
and adds a per-section UoC mapping (ex.uoc "Evidences:" line) so it doubles as the marking
reference. Covers the two 503 design strands: the web-scale architecture (§5) and the
audit-log microservice (§6). Disaster recovery is deliberately NOT here — it is the companion
DR Plan (Part B). Data residency is treated as a fixed input, not a deliverable.

Figures are indicative and internally consistent with the DR Plan exemplar (primary
ap-southeast-2 Sydney; India slice ap-south-1 Mumbai; availability >= 99.9%).

Reuses the docx brand helpers (build_bc_template) + the exemplar helpers (build_bc_exemplar).

Usage:  python scripts/build_at1_solution_design_exemplar.py [output.docx]
Default: S1-CL2-Cloud-Disaster-Recovery/assessments/AT1/AT1-exemplar-solution-design.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_bc_template as bc  # noqa: E402  (shared branding helpers + palette)
import build_bc_exemplar as ex  # noqa: E402  (uoc, para, bullets, etable)

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def build(path):
    doc = Document()
    bc.configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    bc.build_header_footer(sec)

    # ---- COVER ----
    bc.wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(bc.ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(bc.GREY)
    bc.paragraph_bottom_rule(doc.add_paragraph(), bc.TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Solution Design")
    sub = doc.add_paragraph().add_run("YAT LMS Global Expansion — Solution Design")
    sub.font.size = Pt(15); sub.bold = True; sub.font.color.rgb = RGBColor.from_string(bc.TERRACOTTA)
    note = doc.add_paragraph().add_run("Assessor exemplar — internal marking reference (not for distribution to students)")
    note.italic = True; note.font.size = Pt(10); note.font.color.rgb = RGBColor.from_string(bc.GREY)
    doc.add_paragraph()
    cover = [
        ("Engagement", "YAT LMS Global Expansion & Disaster Recovery"),
        ("Engagement reference", "YAT-MTS-2026-002"),
        ("Document version", "v1.0"),
        ("Prepared by", "MTS Consultant"),
        ("Date", "[DD/MM/YYYY]"),
        ("Submitted to", "Sam Walker (YAT ICT Manager) · Pat Lin (MTS Senior Consultant)"),
        ("Related documents", "Functional & Non-Functional Requirements; Data Residency & Sovereignty "
                              "Requirements; Cloud Baseline Solution Design; companion Disaster Recovery Plan"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        bc.set_cell_borders(cells[0]); bc.set_cell_borders(cells[1]); bc.shade_cell(cells[0], bc.CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    bc.add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    # 1 Purpose and Scope
    h1("1. Purpose and Scope")
    ex.uoc(doc, "[ICTCLD503 PC 1.1] determine and confirm cloud web-scaling needs")
    ex.para(doc, "This Solution Design sets out the architecture changes that prepare the YAT Learning "
                 "Management System to serve the new India-campus cohort alongside the existing Australian "
                 "cohort. Two changes are required: (1) scale the web application to serve a global user base "
                 "while maintaining availability and security, and (2) add a microservice that records "
                 "India-cohort access events to a store held in India, meeting the data-residency "
                 "requirements. This design is the basis for the Phase 2 implementation (the Deployment "
                 "Report) and for the companion Disaster Recovery Plan.")
    ex.para(doc, "In scope: the web-tier scaling design (edge delivery, compute/cache/storage scaling) and "
                 "the audit-log microservice design. Out of scope: disaster recovery (the companion DR Plan), "
                 "the build and its evidence (the Phase 2 Deployment Report), and the legal determination of "
                 "the residency obligations (owned by YAT compliance — this design treats them as fixed "
                 "inputs).")

    # 2 Design Inputs and Requirements
    h1("2. Design Inputs and Requirements")
    h3("2.1 Inputs")
    ex.uoc(doc, "[ICTCLD503 PC 1.1] confirm web-scaling needs according to business needs")
    ex.para(doc, "This design is built from the Functional & Non-Functional Requirements, the Data Residency & "
                 "Sovereignty Requirements, the LMS Application Specification (cloud), and the existing Cloud "
                 "Baseline Solution Design. The partnership roughly doubles the active user base and adds a "
                 "cohort about 7,000 km from the Sydney region, so the application must serve India users with "
                 "acceptable latency, absorb a higher and more variable concurrent load, and preserve the "
                 "existing availability and security posture.")
    h3("2.2 Requirements the design must meet")
    ex.etable(doc, ["Ref", "Requirement", "Source"],
              [["R1", "Serve a global user base (AU + India) with acceptable latency", "Functional & Non-Functional Requirements"],
               ["R2", "Scale network, compute and storage as utilisation increases", "Functional & Non-Functional Requirements"],
               ["R3", "Maintain availability (>= 99.9%) and security through the changes", "Requirements / Baseline Design"],
               ["R4", "India-cohort access logs retained in India (>= 180 days)", "Data Residency Requirements (CERT-In)"],
               ["R5", "Main LMS data may remain in Australia (light residency slice)", "Data Residency Requirements (DPDP)"]],
              widths=[1.4, 8.6, 6.0])

    # 3 Review of the Existing Architecture
    h1("3. Review of the Existing Architecture")
    ex.uoc(doc, "[ICTCLD503 PC 1.2] review architecture for web application according to business needs")
    ex.para(doc, "The current architecture (Cloud Baseline Solution Design) is a single-region, multi-AZ "
                 "deployment in ap-southeast-2 (Sydney): an Application Load Balancer across two Availability "
                 "Zones, an Auto Scaling group of EC2 application instances, a Multi-AZ RDS database, and S3 "
                 "for static assets. It is highly available within the region but was sized for a single "
                 "Australian cohort served entirely from Sydney.")
    ex.para(doc, "Reviewed layer by layer against R1–R3, three gaps drive this design:")
    ex.bullets(doc, [
        "Edge / global delivery — India users are served entirely from Sydney; there is no CDN, so static "
        "and cacheable content travels the full distance on every request, hurting latency (R1).",
        "Read-path scale — the Auto Scaling group scales compute, but there is no caching layer to absorb a "
        "larger, spikier read and session load off the database (R2).",
        "Logging residency — all access logging currently lands in Sydney; there is no mechanism to keep "
        "India-cohort access logs in India (R4).",
    ])
    ex.para(doc, "The compute and load-balancing tiers scale adequately and are retained largely unchanged; "
                 "the gaps are at the edge, the data read path, and logging residency — addressed in §5 and §6.")

    # 4 Architecture Design
    h1("4. Architecture Design")
    ex.uoc(doc, "[ICTCLD503 PC 1.4] design architecture changes using cloud services (overview) · [ICTCLD503 PC 1.5] scale for a global user base")
    h3("4.1 Assumptions and constraints")
    ex.bullets(doc, [
        "CL1 is fully implemented: the LMS runs highly available in ap-southeast-2 (multi-AZ ALB/ASG/RDS).",
        "The change is incremental — this design modifies the baseline rather than rebuilding it.",
        "Data residency is a fixed input: India-cohort access logs must be in India; the main LMS database "
        "may remain in Australia (the light residency slice).",
    ])
    h3("4.2 AWS accounts and regions")
    ex.para(doc, "ap-southeast-2 (Sydney) remains the primary application and data region. ap-south-1 (Mumbai) "
                 "is introduced for one purpose only — to hold the India-cohort audit log (the residency "
                 "slice, §6). Global users are reached through edge delivery (§5.3), not by replicating the "
                 "application into other regions. The main LMS database stays in Sydney; only the access logs "
                 "are residency-bound.")
    h3("4.3 Design relative to the baseline")
    ex.para(doc, "The baseline's IAM, VPC / network, compute (ALB + Auto Scaling group), database (Multi-AZ "
                 "RDS) and S3 storage are retained. This design adds three things: a global edge / caching "
                 "layer in front of the application (§5), a read / session cache (§5), and the audit-log "
                 "microservice in India (§6). Infrastructure layers not described there are unchanged from the "
                 "Baseline Solution Design.")

    # 5 Web-scale Design
    h1("5. Web-scale Design")
    ex.uoc(doc, "[ICTCLD503 PC 1.3, 1.4, 1.5, 1.6, 1.7] · [ICTCLD503 PE 1, 5] · [ICTCLD503 KE 3, 6]")
    h3("5.1 Web-scaling needs")
    ex.para(doc, "The drivers (PC 1.1, 1.3): roughly double the concurrent users, a geographically distant "
                 "second cohort, and a spikier demand profile (assessment deadlines, term start). The design "
                 "must scale network, compute and storage elastically (R2) and serve the global base with "
                 "acceptable latency (R1) without weakening availability or security (R3).")
    h3("5.2 Scaling by layer")
    ex.para(doc, "Services selected to scale the web application (PC 1.3) and how each layer scales as "
                 "utilisation increases (PC 1.4):")
    ex.etable(doc, ["Layer", "Service", "How it scales"],
              [["Edge / network", "CloudFront CDN + Route 53 latency routing",
                "Serves cached / static content from edge locations worldwide; India users hit the nearest edge, not Sydney"],
               ["Compute", "ALB + EC2 Auto Scaling (retained)",
                "Target-tracking policies add / remove instances on demand; retuned for the higher global load"],
               ["Caching / read path", "ElastiCache (Redis) — new",
                "Absorbs session state and hot reads off the database as concurrency grows"],
               ["Storage", "S3 (retained), served via CloudFront",
                "Effectively unlimited object storage; edge-cached for global delivery"],
               ["Database", "RDS Multi-AZ (retained, Sydney)",
                "Vertical scaling, with a read replica as future headroom (not required at launch)"]],
              widths=[3.4, 5.0, 7.6])
    h3("5.3 Global delivery")
    ex.para(doc, "A global user base is served (PC 1.5) by pushing content to the edge with CloudFront and "
                 "routing users to the lowest-latency entry point with Route 53. Dynamic requests still reach "
                 "the Sydney origin, but the dominant share of LMS traffic — course media, static assets, "
                 "cacheable GETs — is served from the nearest edge, so India users get acceptable latency "
                 "without replicating the application or its data into India.")
    h3("5.4 Web-scale component choices")
    ex.para(doc, "Each web-scale component is chosen where it fits (KE 3):")
    ex.etable(doc, ["Choice", "Decision", "Why"],
              [["SQL vs NoSQL", "SQL (RDS) for LMS data; NoSQL (DynamoDB) for the audit log",
                "Relational integrity for enrolment / assessment data; append-only high-write key-value for logs (§6)"],
               ["Monolith vs microservice", "Keep the LMS a monolith; split the audit function into a microservice",
                "The monolith performs well and isn't the bottleneck; the new, independently-scaling concern is isolated (§6)"],
               ["Compute model", "Container / VM (EC2 + ASG) for the app; serverless (Lambda) for the microservice",
                "Steady app load suits always-on instances; bursty event load suits scale-to-zero serverless"],
               ["CDN + in-memory store", "CloudFront (CDN) + ElastiCache (in-memory)",
                "Edge delivery for global latency; in-memory caching for read scale"]],
              widths=[3.6, 5.8, 6.6])
    h3("5.5 Availability and security maintained")
    ex.para(doc, "Availability is preserved (PC 1.6) — the multi-AZ ALB / ASG / RDS design is unchanged, and "
                 "CloudFront and ElastiCache are themselves resilient managed services. Security is preserved "
                 "— TLS end-to-end, AWS WAF at the CloudFront edge, the origin locked to accept only "
                 "CloudFront traffic, and no change to the IAM or data-at-rest posture. The design was "
                 "reviewed against R1–R3 and meets them.")
    ex.para(doc, "This is one architecture that scales networking, compute and storage for the multi-tier LMS "
                 "(PE 1) using established web-scaling principles — edge offload, horizontal compute scaling, "
                 "read caching (PE 5, KE 6) — and the rationale for each change is recorded above, satisfying "
                 "the requirement to document and justify the architecture changes (PC 1.7).")

    # 6 Microservice and Serverless Design
    h1("6. Microservice and Serverless Design")
    ex.uoc(doc, "[ICTCLD503 PC 2.1, 2.2, 2.3, 2.4] · [ICTCLD503 PE 2] · [ICTCLD503 KE 4]")
    h3("6.1 Microservices and data transactions")
    ex.para(doc, "One microservice is introduced (PC 2.1) — the Access-Log Service. Its single responsibility "
                 "is to record India-cohort access events durably and in-region. The data transaction: on each "
                 "relevant LMS event (login, course access, assessment view) for an India-cohort user, the LMS "
                 "emits an event and the service persists one append-only audit record. It never reads back "
                 "into the LMS and holds no other state — highly cohesive and loosely coupled.")
    h3("6.2 Supporting cloud services")
    ex.para(doc, "The service is serverless and event-driven, hosted in ap-south-1 (India) (PC 2.2):")
    ex.etable(doc, ["Concern", "Service", "Why"],
              [["Ingress / API", "API Gateway (regional, ap-south-1)", "Receives the LMS event webhook over HTTPS"],
               ["Messaging / queuing", "SQS queue", "Buffers events so a spike or downstream slowness never blocks or loses one"],
               ["Compute", "AWS Lambda", "Serverless; scales to zero and to spikes, with no servers to manage"],
               ["Persistent storage", "DynamoDB (ap-south-1)", "Append-only NoSQL audit log, kept in India (R4 / CERT-In >= 180 days)"]],
              widths=[3.6, 5.0, 7.4])
    h3("6.3 Microservice architecture")
    ex.para(doc, "Flow (PC 2.3): LMS event webhook → API Gateway (ap-south-1) → SQS → Lambda → DynamoDB "
                 "(ap-south-1). The LMS is the event producer and knows only the webhook contract; everything "
                 "behind it can change independently. The design is highly cohesive and loosely coupled (the "
                 "service does one thing; the LMS integrates only through the contract); it uses a managed "
                 "database / storage service for persistent data (DynamoDB); and it connects its parts with "
                 "API, messaging and queuing services (API Gateway + SQS) — the three KE 4 elements. Because "
                 "the store is in ap-south-1, the access logs are resident in India by construction.")
    h3("6.4 Interface / integration contract")
    ex.para(doc, "The contract the LMS calls is the single coupling point (PC 2.4): POST /events with a JSON "
                 "body { event_id, occurred_at, user_ref, cohort, event_type, source_ip }, authenticated with "
                 "a signing key. Defining it here lets the LMS and the service evolve independently. The same "
                 "generic contract is reused unchanged by the Accounting practice project, where the payload "
                 "carries transaction-audit events instead of access events — the engine is identical, only "
                 "the payload differs.")
    ex.para(doc, "This is one microservice architecture for a simple application (PE 2), designed on serverless "
                 "cloud services, with each choice justified above — satisfying the requirement to document "
                 "and justify the microservice design (PC 2.4).")

    # 7 Implementation Sequencing
    h1("7. Implementation Sequencing")
    ex.para(doc, "The changes are applied in an order that holds availability throughout; the build and its "
                 "evidence are the Phase 2 Deployment Report.")
    ex.etable(doc, ["#", "Change", "Expected impact"],
              [["1", "Stand up the audit-log microservice in ap-south-1 (API Gateway → SQS → Lambda → DynamoDB) and point the LMS webhook at it", "None — additive"],
               ["2", "Introduce ElastiCache and move the session / read path onto it", "None — cache warms transparently"],
               ["3", "Put CloudFront + WAF in front of the ALB and switch Route 53 to the distribution", "None — origin unchanged; cutover at DNS"],
               ["4", "Retune the Auto Scaling group scaling policies for the global load profile", "None"]],
              widths=[0.9, 10.6, 4.5])

    # 8 Simulation and Verification Plan
    h1("8. Simulation and Verification Plan")
    ex.para(doc, "How the design will be shown to work; the evidence is produced in Phase 2.")
    h3("8.1 Web-scale verification")
    ex.bullets(doc, [
        "A load test ramps concurrent users to the projected global peak; the Auto Scaling group scales out "
        "and back, edge offload and cache hit-rate are measured, and p95 latency for an India-region client "
        "stays within target (R1, R2).",
        "Confirm multi-AZ failover still works, WAF blocks the OWASP baseline, and the origin rejects "
        "non-CloudFront traffic (R3).",
    ])
    h3("8.2 Microservice verification")
    ex.bullets(doc, [
        "Fire sample webhook events; confirm each is persisted exactly once in the ap-south-1 DynamoDB table.",
        "Force a downstream stall; confirm events queue in SQS and drain without loss.",
        "Confirm the audit store location is ap-south-1 (R4).",
    ])

    # 9 Out of Scope
    h1("9. Out of Scope")
    ex.bullets(doc, [
        "Disaster recovery — covered by the companion DR Plan (Part B).",
        "The implementation and its evidence — the Phase 2 Deployment Report (AT2).",
        "The legal determination of the residency obligations — owned by YAT compliance; treated here as "
        "fixed inputs.",
        "Replicating the main LMS database into India — deliberately avoided; the residency slice is the "
        "access logs only, and replication would add cost and complexity for no required benefit.",
    ])

    # 10 References
    h1("10. References")
    ex.bullets(doc, [
        "Functional & Non-Functional Requirements (YAT LMS Global Expansion)",
        "Data Residency & Sovereignty Requirements (DPDP Act 2023; CERT-In Directions 2022)",
        "LMS Application Specification (cloud)",
        "Cloud Baseline Solution Design",
        "AWS Well-Architected Framework — Performance Efficiency and Reliability pillars",
    ])

    # 11 Review and Approval
    h1("11. Review and Approval")
    ex.uoc(doc, "[ICTCLD503 PC 1.6] review design as required")
    ex.para(doc, "This design was reviewed by the MTS Senior Consultant and presented to the YAT ICT Manager "
                 "for approval at the design walkthrough (the AT1 presentation), together with the companion "
                 "DR Plan. Approval at that walkthrough is the gate to Phase 2 implementation.")
    h3("11.1 Reviewer feedback and author response")
    ex.etable(doc, ["#", "Reviewer feedback", "Author response", "Resulting change"],
              [["1", "Confirm India users are served without replicating data into India",
                "Edge delivery (CloudFront) serves the global base; only the access logs sit in India",
                "Clarified in §4.2 and §5.3"],
               ["2", "Confirm the microservice cannot lose events under load",
                "An SQS buffer decouples ingress from the writer, so spikes queue rather than drop",
                "Recorded in §6.2 / §6.3"]],
              widths=[0.8, 5.6, 5.4, 4.2])
    h3("11.2 Sign-off")
    ex.etable(doc, ["Role", "Name", "Decision", "Date"],
              [["Prepared by", "MTS Consultant", "Submitted", "[date]"],
               ["Reviewed by", "Pat Lin (MTS Senior Consultant)", "Approved for submission", "[date]"],
               ["Approved by", "Sam Walker (YAT ICT Manager)", "Approved at design walkthrough — cleared to implement", "[date]"]],
              widths=[3.6, 5.0, 5.4, 2.0])

    # Document control
    h1("Document control")
    ex.etable(doc, ["Field", "Value"],
              [["Document version", "v1.0 — Solution Design"],
               ["Author", "MTS Consultant"],
               ["Engagement", "YAT LMS Global Expansion & Disaster Recovery"],
               ["Companion document", "Disaster Recovery Plan"],
               ["Implemented by", "LMS Global Expansion Deployment Report (Phase 2)"]],
              widths=[5.0, 10.5])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL2-Cloud-Disaster-Recovery/assessments/AT1/AT1-exemplar-solution-design.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
