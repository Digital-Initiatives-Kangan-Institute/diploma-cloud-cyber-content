"""One-off fix for voice slips + stray brackets in AT1-BusinessCase-Student.docx.

Applies targeted text replacements at the paragraph level. For each replacement,
finds the spanning runs and rewrites the first one, blanking out runs after it
in the same span. Preserves formatting on the first run of each matched span and
on all unaffected runs.
"""

import shutil
import sys
from pathlib import Path

from docx import Document

DOCX = Path(
    "c:/Users/micro/Documents/Kangan/diploma-cloud-cyber/"
    "S1-CL1-Cloud-Design-Build/assessments/AT1/AT1-BusinessCase-Student.docx"
)

REPLACEMENTS = [
    ("allowing the you to present", "allowing you to present"),
    ("guide the you to specific recommendations", "guide you to specific recommendations"),
    ("coach the you through your responses", "coach you through your responses"),
    ("declaration that all work is their own", "declaration that all work is your own"),
    ("for this assessment the you must", "for this assessment you must"),
    ("in their own words (not verbatim", "in your own words (not verbatim"),
    ("specific sections of their own Business Case", "specific sections of your own Business Case"),
    ("You seek feedback during Q&A and responds substantively",
     "You seek feedback during Q&A and respond substantively"),
    ("You use plain English and translates technical terminology",
     "You use plain English and translate technical terminology"),
    # Stray brackets in criteria table
    ("difficulty assessment for both options]", "difficulty assessment for both options"),
    ("deferred to later sign-off gates [", "deferred to later sign-off gates"),
]


def fix_paragraph(paragraph, find: str, replace: str) -> bool:
    """If `find` appears in paragraph.text, rewrite the spanning runs.

    Strategy: walk runs accumulating text; once the accumulated text crosses the
    start of `find`, the run that contained that boundary owns the replacement.
    Subsequent runs covering the rest of `find` get their text cleared.

    Returns True if a replacement was made.
    """
    full_text = paragraph.text
    if find not in full_text:
        return False

    start = full_text.index(find)
    end = start + len(find)

    runs = paragraph.runs
    offset = 0
    new_run_texts = []
    span_start_run = None
    span_end_run = None
    span_start_offset = None
    span_end_offset = None

    for i, r in enumerate(runs):
        r_start = offset
        r_end = offset + len(r.text)
        new_run_texts.append(r.text)

        if span_start_run is None and r_end > start:
            span_start_run = i
            span_start_offset = start - r_start
        if span_end_run is None and r_end >= end:
            span_end_run = i
            span_end_offset = end - r_start

        offset = r_end

    if span_start_run is None or span_end_run is None:
        return False

    if span_start_run == span_end_run:
        # All within one run — simple replace inside that run's text
        original = new_run_texts[span_start_run]
        new_run_texts[span_start_run] = (
            original[:span_start_offset] + replace + original[span_end_offset:]
        )
    else:
        # Spans multiple runs. Stitch: first run keeps prefix + replacement;
        # middle runs get cleared; last run keeps suffix.
        start_original = new_run_texts[span_start_run]
        end_original = new_run_texts[span_end_run]
        new_run_texts[span_start_run] = start_original[:span_start_offset] + replace
        for i in range(span_start_run + 1, span_end_run):
            new_run_texts[i] = ""
        new_run_texts[span_end_run] = end_original[span_end_offset:]

    for r, new_text in zip(runs, new_run_texts):
        r.text = new_text

    return True


def main():
    backup = DOCX.with_suffix(".docx.bak")
    shutil.copy2(DOCX, backup)
    print(f"Backup written to {backup}")

    doc = Document(str(DOCX))

    def iter_paragraphs(d):
        yield from d.paragraphs
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
                    for nested in cell.tables:
                        for nrow in nested.rows:
                            for ncell in nrow.cells:
                                yield from ncell.paragraphs

    paragraphs = list(iter_paragraphs(doc))

    results = []
    for find, replace in REPLACEMENTS:
        hits = 0
        for p in paragraphs:
            if fix_paragraph(p, find, replace):
                hits += 1
        results.append((find, replace, hits))

    doc.save(str(DOCX))

    print()
    print(f"{'hits':>5}  find -> replace")
    print("-" * 70)
    any_zero = False
    for find, replace, hits in results:
        marker = "  " if hits else "!!"
        if not hits:
            any_zero = True
        print(f"{marker}{hits:>3}  {find!r}  ->  {replace!r}")

    if any_zero:
        print("\nSome targets had 0 hits — investigate before trusting the output.")
        sys.exit(1)
    print("\nAll targets replaced.")


if __name__ == "__main__":
    main()
